import gradio as gr
import fitz  # PyMuPDF
import re
from together import Together
import io
import os
from dotenv import load_dotenv
from docx import Document

load_dotenv()  # load environment variables from .env
# Get the absolute path of the current script
current_dir = os.path.dirname(os.path.abspath(__file__))

sys_pro_main = os.path.join(current_dir, 'system_prompt_main.txt')
sys_summary_path = os.path.join(current_dir, 'system_prompt_summary.txt')
coherence_path = os.path.join(current_dir, 'prompt_coherence.txt') 

API_key_1 = os.getenv("TOGETHER_API_KEY_1")
API_key_2 = os.getenv("TOGETHER_API_KEY_2")
API_key_3 = os.getenv("TOGETHER_API_KEY_3")

if not API_key_1 or not API_key_2 or not API_key_3:
    raise ValueError("All three API keys must be set in environment variables")

with open(sys_pro_main, 'r', encoding='utf-8') as file:
    sys_prompt_main = file.read()
with open(sys_summary_path, 'r', encoding='utf-8') as file:
    sys_prompt_summary = file.read()
with open(coherence_path, 'r', encoding='utf-8') as file:
    sys_coherence = file.read()

client_1 = Together(api_key=API_key_1)
client_2 = Together(api_key=API_key_2)
client_3 = Together(api_key=API_key_3)

def extract_sections_from_text(text):
    sections = [
        "TITLE",
        "FIELD OF INVENTION",
        "TECHNICAL FIELD",
        "BACKGROUND",
        "PRIOR ART",
        "OBJECT OF INVENTION",
        "SUMMARY",
        "SUMMARY OF THE INVENTION",
        "BRIEF DESCRIPTION OF FIGURES",
        "DETAILED DESCRIPTION",
        "DETAILED DESCRIPTION OF INVENTION",
        "CLAIMS",
        "ABSTRACT"
    ]
    pattern = r'(?i)^\s*(' + '|'.join(sections) + r')\s*$'
    parts = re.split(pattern, text, flags=re.MULTILINE)
    extracted = {}
    for i in range(1, len(parts), 2):
        header = parts[i].strip().upper()
        content = parts[i+1].strip()
        if header in extracted:
            extracted[header] += "\n\n" + content
        else:
            extracted[header] = content
    return extracted

# def save_outputs(folder_name, output1, output2):
#     if not folder_name:
#         return "Folder name cannot be empty."
#     base_dir = "C:\\Users\\meetm\\Downloads\\SUMMER\\pdf_code\\saved"
#     save_path = os.path.join(base_dir, folder_name)
#     os.makedirs(save_path, exist_ok=True)
#     try:
#         doc1 = Document()
#         doc1.add_heading("Together AI Response - Main", level=1)
#         doc1.add_paragraph(output1)
#         doc1.save(os.path.join(save_path, "main_doc.docx"))

#         doc2 = Document()
#         doc2.add_heading("Together AI Response - Summary", level=1)
#         doc2.add_paragraph(output2)
#         doc2.save(os.path.join(save_path, "summary.docx"))

#         return f"Outputs saved successfully in folder: {save_path}"
#     except Exception as e:
#         return f"Error saving files: {e}"

def extract_pdf_sections(file):
    doc = fitz.open(file.name)
    full_text = ""
    for page in doc:
        full_text += page.get_text()
    sections = extract_sections_from_text(full_text)
    if not sections:
        return "No recognizable patent sections found."
    extracted_text = ""
    for sec, content in sections.items():
        extracted_text += f"=== {sec} ===\n{content}\n\n"
    return extracted_text

def call_together_llm(text):
    output_stream = io.StringIO()
    try:
        response = client_1.chat.completions.create(
            model="Qwen/Qwen2.5-72B-Instruct-Turbo",
            messages=[
                {"role": "system", "content": sys_prompt_main},
                {"role": "user", "content": text}
            ],
            temperature=0.01,
            stream=True
        )
        for token in response:
            if hasattr(token, 'choices') and token.choices[0].delta.content:
                output_stream.write(token.choices[0].delta.content)
    except Exception as e:
        output_stream.write(f"[Error calling Together API: {e}]")
    return output_stream.getvalue()

def call_together_llm_v2(text):
    output_stream = io.StringIO()
    try:
        response = client_2.chat.completions.create(
            model="Qwen/Qwen2.5-72B-Instruct-Turbo",
            messages=[
                {"role": "system", "content": sys_prompt_summary},
                {"role": "user", "content": text}
            ],
            temperature=0.01,
            stream=True
        )
        for token in response:
            if hasattr(token, 'choices') and token.choices[0].delta.content:
                output_stream.write(token.choices[0].delta.content)
    except Exception as e:
        output_stream.write(f"[Error calling Together API: {e}]")
    return output_stream.getvalue()

def save_outputs(
    folder_name,
    main_text,       # llm_output_box
    summary_text,    # llm_output_box_v2
    problem_text,    # problem_box
    summary_abs_text,# summary_box
    features_text,   # features_box
    prototype_text,  # prototype_box
    impact_text,     # impact_box
    applications_text,# applications_box
    keywords_text,   # keywords_box
    figures_text     # figures_box
):
    if not folder_name:
        return "Folder name cannot be empty."

    sys_pro_main = os.path.join(current_dir, 'system_prompt_main.txt')
    base_dir = os.path.join(current_dir,  'document_generated')
    # base_dir = "C:\\Users\\meetm\\Downloads\\SUMMER\\pdf_code\\saved"
    save_path = os.path.join(base_dir, folder_name)
    os.makedirs(save_path, exist_ok=True)

    try:
        # 1) Save main_doc.docx
        doc1 = Document()
        doc1.add_heading("Together AI Response - Main", level=1)
        doc1.add_paragraph(main_text)
        doc1.save(os.path.join(save_path, "main_doc.docx"))

        # 2) Save summary.docx
        doc2 = Document()
        doc2.add_heading("Together AI Response - Summary", level=1)
        doc2.add_paragraph(summary_text)
        doc2.save(os.path.join(save_path, "summary.docx"))

        # 3) Save each split‐section as <Heading>.txt
        sections = {
            "Problem Statement": problem_text,
            "Summary Abstract": summary_abs_text,
            "Salient Features": features_text,
            "Prototype Details": prototype_text,
            "Societal Impact": impact_text,
            "Applications": applications_text,
            "Keywords": keywords_text,
            "Figures": figures_text
        }
        for heading, content in sections.items():
            if content.strip():
                file_name = f"{heading}.txt"
                with open(os.path.join(save_path, file_name), "w", encoding="utf-8") as f:
                    f.write(content)

        return f"Outputs saved successfully in folder: {save_path}"
    except Exception as e:
        return f"Error saving files: {e}"


def call_together_llm_v3(text, user_info_box):
    output_stream = io.StringIO()
    try:
        response = client_3.chat.completions.create(
            model="Qwen/Qwen2.5-72B-Instruct-Turbo",
            messages=[
                {"role": "system", "content": sys_coherence + text},
                {"role": "user", "content": user_info_box}
            ],
            temperature=0.01,
            stream=True
        )
        for token in response:
            if hasattr(token, 'choices') and token.choices[0].delta.content:
                output_stream.write(token.choices[0].delta.content)
    except Exception as e:
        output_stream.write(f"[Error calling Together API: {e}]")
    return output_stream.getvalue()

# ------------ NEW FUNCTION TO SPLIT “@#$% Heading:” RESPONSES ------------
def split_response(raw_text):
    """
    Parses text of the form "@#$% Heading: <content>" and returns a tuple
    of eight strings (one for each expected section). If a section is missing,
    returns an empty string for that slot.
    """
    pattern = r'@#\$%\s*(.+?):\s*([\s\S]*?)(?=(?:@#\$%|\Z))'
    matches = re.findall(pattern, raw_text)
    sections = { heading.strip(): content.strip() for heading, content in matches }

    return (
        sections.get("Problem Statement", ""),
        sections.get("Summary Abstract", ""),
        sections.get("salient_features", ""),
        sections.get("prototype_details", ""),
        sections.get("societal_impact", ""),
        sections.get("applications", ""),
        sections.get("keywords", ""),
        sections.get("Figures", "")
    )
# -----------------------------------------------------------------------

# with gr.Blocks() as iface:
#     gr.Markdown("### Patent Extractor with Together AI Integration")

#     # --- Step 1 / Step 2 / Step 3 / Step 4 UI as before ---
#     file_input = gr.File(file_types=[".pdf"], label="Upload Indian Patent PDF")
#     extracted_textbox = gr.Textbox(lines=20, label="Extracted Patent Sections")
#     llm_output_box = gr.Textbox(lines=15, label="Together AI Response")
#     llm_output_box_v2 = gr.Textbox(lines=10, label="Together AI Response (V2)")
#     llm_output_box_v3 = gr.Textbox(lines=10, label="Together AI Response (Coherence)")

#     folder_name_input = gr.Textbox(label="Enter folder name to save outputs")
#     save_button = gr.Button("Save Outputs to Folder")
#     save_status = gr.Textbox(label="Save Status", interactive=False)

#     extract_button = gr.Button("Step 1: Extract Text")
#     send_button = gr.Button("Step 2: Send to Together AI")
#     send_button_v2 = gr.Button("Step 3: Send to Together AI (V2)")
#     send_button_v3 = gr.Button("Step 4: Send to Together AI (Coherence)")

#     extract_button.click(fn=extract_pdf_sections, inputs=file_input, outputs=extracted_textbox)
#     send_button.click(fn=call_together_llm, inputs=extracted_textbox, outputs=llm_output_box)
#     send_button_v2.click(fn=call_together_llm_v2, inputs=extracted_textbox, outputs=llm_output_box_v2)
#     send_button_v3.click(fn=call_together_llm_v3, inputs=(extracted_textbox, llm_output_box), outputs=llm_output_box_v3)
#     save_button.click(
#         fn=save_outputs,
#         inputs=[
#             folder_name_input,
#             llm_output_box,        # main_text
#             llm_output_box_v2,     # summary_text
#             problem_box,           # problem_text
#             summary_box,           # summary_abs_text
#             features_box,          # features_text
#             prototype_box,         # prototype_text
#             impact_box,            # impact_text
#             applications_box,      # applications_text
#             keywords_box,          # keywords_text
#             figures_box            # figures_text
#         ],
#         outputs=save_status
#     )


#     # --- NEW “Split Sections” UI COMPONENTS ---
#     gr.Markdown("### Parsed Sections (from the '@#$%' response)")

#     problem_box       = gr.Textbox(lines=5, label="Problem Statement")
#     summary_box       = gr.Textbox(lines=5, label="Summary Abstract")
#     features_box      = gr.Textbox(lines=5, label="Salient Features")
#     prototype_box     = gr.Textbox(lines=5, label="Prototype Details")
#     impact_box        = gr.Textbox(lines=5, label="Societal Impact")
#     applications_box  = gr.Textbox(lines=5, label="Applications")
#     keywords_box      = gr.Textbox(lines=5, label="Keywords")
#     figures_box       = gr.Textbox(lines=5, label="Figures")

#     split_button = gr.Button("Split '@#$%' Sections")
#     split_button.click(
#         fn=split_response,
#         inputs=llm_output_box,
#         outputs=[
#             problem_box,
#             summary_box,
#             features_box,
#             prototype_box,
#             impact_box,
#             applications_box,
#             keywords_box,
#             figures_box
#         ]
#     )

# iface.launch()
with gr.Blocks() as iface:
    gr.Markdown("### Patent Extractor with Together AI Integration")

    file_input = gr.File(file_types=[".pdf"], label="Upload Indian Patent PDF")
    extracted_textbox    = gr.Textbox(lines=20, label="Extracted Patent Sections")
    llm_output_box       = gr.Textbox(lines=15, label="Together AI Response")
    llm_output_box_v2    = gr.Textbox(lines=10, label="Together AI Response (V2)")
    llm_output_box_v3    = gr.Textbox(lines=10, label="Together AI Response (Coherence)")

    folder_name_input = gr.Textbox(label="Enter folder name to save outputs")
    save_button       = gr.Button("Save Outputs to Folder")
    save_status       = gr.Textbox(label="Save Status", interactive=False)

    extract_button   = gr.Button("Step 1: Extract Text")
    send_button      = gr.Button("Step 2: Send to Together AI")
    send_button_v2   = gr.Button("Step 3: Send to Together AI (summary)")
    send_button_v3   = gr.Button("Step 4: Send to Together AI (Coherence)")

    extract_button.click(fn=extract_pdf_sections, inputs=file_input, outputs=extracted_textbox)
    send_button.click(fn=call_together_llm, inputs=extracted_textbox, outputs=llm_output_box)
    send_button_v2.click(fn=call_together_llm_v2, inputs=extracted_textbox, outputs=llm_output_box_v2)
    send_button_v3.click(fn=call_together_llm_v3, inputs=(extracted_textbox, llm_output_box), outputs=llm_output_box_v3)

    # # ─── Declare the eight split‐section textboxes before using them ───
    # gr.Markdown("### Parsed Sections (from the '@#$%' response)")
    # problem_box       = gr.Textbox(lines=5, label="Problem Statement")
    # summary_box       = gr.Textbox(lines=5, label="Summary Abstract")
    # features_box      = gr.Textbox(lines=5, label="Salient Features")
    # prototype_box     = gr.Textbox(lines=5, label="Prototype Details")
    # impact_box        = gr.Textbox(lines=5, label="Societal Impact")
    # applications_box  = gr.Textbox(lines=5, label="Applications")
    # keywords_box      = gr.Textbox(lines=5, label="Keywords")
    # figures_box       = gr.Textbox(lines=5, label="Figures")

    # split_button = gr.Button("Split '@#$%' Sections")
    # split_button.click(
    #     fn=split_response,
    #     inputs=llm_output_box,
    #     outputs=[
    #         problem_box,
    #         summary_box,
    #         features_box,
    #         prototype_box,
    #         impact_box,
    #         applications_box,
    #         keywords_box,
    #         figures_box
    #     ]
    # )
    # ────────────────────────────────────────────────────────────────────

    # ─── Now wire up Save Outputs with all ten inputs ───
    save_button.click(
        fn=save_outputs,
        inputs=[
            folder_name_input,   # 1
            llm_output_box,      # 2: main_text
            llm_output_box_v2   # 3: summary_text
            # problem_box,         # 4: problem_text
            # summary_box,         # 5: summary_abs_text
            # features_box,        # 6: features_text
            # prototype_box,       # 7: prototype_text
            # impact_box,          # 8: impact_text
            # applications_box,    # 9: applications_text
            # keywords_box,        #10: keywords_text
            # figures_box          #11: figures_text
        ],
        outputs=save_status
    )
    # ────────────────────────────────────────────────────────────────

iface.launch()
