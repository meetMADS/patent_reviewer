from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_HEADER_FOOTER
from docx.shared import Inches, Pt, RGBColor, Cm
from docxcompose.composer import Composer
import os
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.oauth2 import service_account
from googleapiclient.http import MediaIoBaseUpload

main_doc = Document()

app = Flask(__name__)

# Google Drive API setup
current_dir = os.path.dirname(os.path.abspath(__file__))
SCOPES = ['https://www.googleapis.com/auth/drive']
SERVICE_ACCOUNT_FILE = os.path.join(current_dir, 'ircc-462113-96ae13d4fa40.json')
logo_path = os.path.join(current_dir, 'new.png') 
app_path = os.path.join(current_dir, 'appendix.docx') 

@app.route('/')
def form():
    return render_template('form.html')
def index():
    return render_template('form.html')
def create_drive_service():
    creds = service_account.Credentials.from_service_account_file(
        SERVICE_ACCOUNT_FILE, scopes=SCOPES)
    drive_service = build('drive', 'v3', credentials=creds)
    return drive_service

def upload_file_to_drive(drive_service, file_name, parent_folder_id):
    file_metadata = {
        'name': file_name,
        'parents': [parent_folder_id]
    }
    media = MediaFileUpload(file_name, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id'
    ).execute()
    print(f"File '{file_name}' uploaded to Google Drive. File ID: {file.get('id')}")
    return file.get('id')

@app.route('/fetch_txt', methods=['POST'])
def fetch_txt():
    data = request.get_json()
    folder_path = data.get('folder_path')
    txt_file = next((f for f in os.listdir(folder_path) if f.endswith('.txt')), None)
    
    result = {}
    if txt_file:
        with open(os.path.join(folder_path, txt_file), 'r', encoding='utf-8') as f:
            for line in f:
                if ':' in line:
                    key, value = line.split(':', 1)
                    normalized = key.strip().lower().replace(" ", "_")
                    
                    # ** Map “idf_no” → “idf” so it matches <input name="idf"> **
                    if normalized == "idf_no":
                        normalized = "idf"
                    # If you have other mismatches, do the same:
                    # e.g. map “type_of_ip” → “ip_type”, etc.
                    elif normalized == "type_of_ip":
                        normalized = "ip_type"
                    elif normalized == "related_patents_country":
                        normalized = "related_patent_country"
                    # …and so on for any other field whose “normalized” string doesn’t match your form’s name…
                    
                    result[normalized] = value.strip()
    return jsonify(result)


@app.route('/generate', methods=['POST'])
def generate():
    data = request.form.to_dict()  # ✅ Now it's mutable
    data['category'] = request.form.getlist('category')  # ✅ Safe to assign


    category_str = ", ".join(data['category'])  # e.g., "Software, AI-ML, Robotics"
    data['technology_type'] = request.form.getlist('technology_type')
    technology_type_str = ", ".join(data['technology_type'])
    doc = Document()

    def add_table_row(label, value, label_color=None, value_color=None, bold_label=True):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        row = table.rows[0]
        
        # Set label cell text and style
        label_cell = row.cells[0]
        label_cell.text = label

        # Center horizontally (paragraph) and vertically (cell)
        paragraph = label_cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Vertical center: set cell vertical alignment to center
        from docx.enum.table import WD_ALIGN_VERTICAL
        label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Format label font
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.bold = bold_label
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), 'Arial')

        if label_color:
            set_cell_background(label_cell, label_color)

        # Set value cell text and style
        value_cell = row.cells[1]
        value_cell.text = value
        for para in value_cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                rFonts = run._element.rPr.rFonts
                rFonts.set(qn('w:eastAsia'), 'Arial')
        
        if value_color:
            set_cell_background(value_cell, value_color)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Clear all existing paragraphs in header
        for para in header.paragraphs:
            p = para._element
            p.getparent().remove(p)

        # Add a new paragraph and insert image centered
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # paragraph.paragraph_format.left_indent = Inches(-5)
        run = paragraph.add_run()
        run.add_picture( logo_path, width=Inches(7))

    # Set default font to Arial 10 for the Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    # For East Asian fonts (to avoid font reset issues)
    rFonts = font.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Arial')

    # --- Utility functions ---
    def set_cell_background(cell, color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)
    
    def add_section_heading(doc, text, bg_color):
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Merge the two cells into one
        merged_cell = table.cell(0, 0).merge(table.cell(0, 1))
        merged_cell.text = text
        set_cell_background(merged_cell, bg_color)

        paragraph = merged_cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        rFonts = run._element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), 'Arial')

        return table

    def add_horizontal_line(doc):
        p = doc.add_paragraph()
        p_paragraph = p._p  # Access the XML element of the paragraph
        pPr = p_paragraph.get_or_add_pPr()

        # Create the border element
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')        # Thickness
        bottom.set(qn('w:space'), '1')      # Padding
        bottom.set(qn('w:color'), '000000') # Black color

        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section_heading_v2(text, color='000000'):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(color)  # Use hex color like '000000' for black
        
        # Ensure East Asia font is also set for Arial
        rFonts = run._element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), 'Arial')

    def add_bullet_list(items_str, indent=0.75, font_size=10):
        """
        items_str: a newline-separated string (from form textarea).
        Each non-empty line becomes one bullet. If there’s a “:”, text before it is bolded.
        """
        for line in items_str.splitlines():
            text = line.strip()
            if not text:
                continue

            # Create a bulleted paragraph
            p = doc.add_paragraph(style='List Bullet')
            pf = p.paragraph_format
            pf.left_indent = Inches(indent)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1
            pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Set paragraph alignment to Justify

            # If there’s a colon, split into “heading:” and “rest”
            if ':' in text:
                heading, rest = text.split(':', 1)
                run1 = p.add_run(heading + ': ')
                run1.bold = True
                run1.font.name = 'Arial'
                run1.font.size = Pt(font_size)
                run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

                run2 = p.add_run(rest.strip())
                run2.font.name = 'Arial'
                run2.font.size = Pt(font_size)
                run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            else:
                # No colon → entire line is normal text
                run = p.add_run(text)
                run.font.name = 'Arial'
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')


    def add_heading_with_text(doc, heading, text, indent=0.25, font_size=10):
        p = doc.add_paragraph(style='No Spacing')
        pf = p.paragraph_format
        pf.left_indent = Inches(indent)
        pf.space_before = Pt(10)
        pf.space_after = Pt(0)
        pf.line_spacing = 1
        pf.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Set paragraph alignment to Justify

        run1 = p.add_run(heading)
        run1.bold = True
        run1.font.name = 'Arial'
        run1.font.size = Pt(font_size)
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        run2 = p.add_run(text)
        run2.font.name = 'Arial'
        run2.font.size = Pt(font_size)
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        return p

    add_table_row("IDF No", data['idf_no'], label_color="FFFFFF", bold_label=True)
    add_table_row("Type of IP", data['ip_type'], label_color="FFFFFF", bold_label=True)
    add_table_row("Category/Theme", category_str, label_color="FFFFFF", bold_label=True)
    add_table_row("Product/Process/Both", data['product_process'], label_color="FFFFFF", bold_label=True)
    add_table_row("Concerned Faculty", data['faculty'], label_color="FFFFFF", bold_label=True)
    add_table_row("Department", data['department'], label_color="FFFFFF", bold_label=True)

    # --- Section: IPR Details ---
    add_section_heading(doc,"IPR Details" , bg_color= "FFFFFF")
    add_table_row("Title", data['title'].title(), bold_label=True)
    add_table_row("Assignee(s)", data['assignees'], bold_label=True)
    add_table_row("Inventor(s)", data['inventors'], bold_label=True)

    # --- Section: Details of IP ---
    add_section_heading(doc, "Details of IP", bg_color="FFFFFF")
    add_table_row("Application No.", data['application_no'], bold_label=False)
    add_table_row("Country", data['country'], bold_label=False)
    add_table_row("Related Patents", data.get('related_patents', ''), bold_label=False)
    add_table_row("Related Patents Country", data.get('related_patent_country', ''), bold_label=False)
    add_table_row("Patent of Addition/Division/Similar technology", data.get('patent_addition_division', ''), bold_label=False)
    add_table_row("Date of Filing", data['date_of_filing'], bold_label=False)
    add_table_row("Grant No.", data['grant_no'], bold_label=False)
    add_table_row("Date of Grant", data['date_of_grant'], bold_label=False)
    add_table_row("Technology readiness level (TRL)", f"Level: {data['trl']}", bold_label=True)
    add_table_row("Type", technology_type_str, value_color="FFFFFF", bold_label=True)


    # --- Section: Commercialisation ---
    add_section_heading(doc, "Status of Commercialisation", bg_color="FFFFFF")
    add_table_row("Available for licensing", data['available_licensing'], label_color="FFFFFF", bold_label=True)
    add_table_row("Under process", data['under_process'], label_color="FFFFFF", bold_label=True)
    add_table_row("Potential sectors to license", data['sectors'], label_color="FFFFFF", bold_label=True)
    add_table_row("Licensed/given to start-up", data['licensed'], label_color="FFFFFF", bold_label=True)
    add_table_row("Name of Licensee/start-up", data['licensee'], label_color="FFFFFF", bold_label=True)
    add_table_row("Tenure of agreement", data['tenure'], label_color="FFFFFF", bold_label=True)
    add_table_row("Available in the market (date if available)", data['market_availability'], label_color="FFFFFF", bold_label=True)
    add_table_row("Unit selling price (if applicable)", data['price'], label_color="FFFFFF", bold_label=True)


    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


    # --- Section 1 ---
    add_section_heading_v2("1. Brief Description (for 1/2 page booklet  - Max. 350 words total):")

    add_heading_with_text(doc, "Problem Statement: ", data['problem'])
    add_heading_with_text(doc, "Salient technical feature(s) / Uniqueness/Advantages: ", "")
    add_bullet_list( data.get('features_3',''))

    add_heading_with_text(doc, "Current Status of Technology: ", data['status'])
    add_heading_with_text(doc, "Societal Impact: ", data['impact'])
    add_heading_with_text(doc, "Applications/Domains/Relevant Industries: ", "")
    add_bullet_list( data.get('applications',''))

    add_horizontal_line(doc)
#############
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Section 2 ---
    add_section_heading_v2("2. Descriptive Details:")

    add_heading_with_text(doc, "Summary/Abstract: ", data['summary'])
    add_heading_with_text(doc, "Problem Statement: ", data['problem'])
    add_heading_with_text(doc, "Salient technical feature(s) / Uniqueness/Advantages: ", "")
    add_bullet_list( data.get('features',''))

    add_heading_with_text(doc, "Prototype Details: ", data['prototype'])
    add_heading_with_text(doc, "Current Status of Technology: ", data['status'])
    add_heading_with_text(doc, "Societal Impact: ", data['impact'])
    add_heading_with_text(doc, "Applications/Domains/Relevant Industries: ", "")
    add_bullet_list( data.get('applications',''))
    add_heading_with_text(doc, "Keywords: ", data['keyword'])
    add_heading_with_text(doc, "Photographs: ", " ")

    doc.add_paragraph("")  # Blank paragraph to add vertical space

    p = doc.add_paragraph()
    p.add_run("Prepared by: " + data['prepared_by']).add_break()
    p.add_run("Faculty: " + data['faculty']).add_break()
    p.add_run("Faculty Signature:").add_break()
    p.add_run("Date: " + data['date'])

    doc.add_paragraph("Note: The write-up may cover aspects such as the problem statement addressed by the technology/its uniqueness/advantages/applications/and key features. Any relevant images or photographs showcasing the technology. This information will be used for promotion and dissemination through the website/internal reports, etc.")

    # Insert a page break before appending appendix
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # Merge the current document with the appendix
    composer = Composer(doc)
    appendix = Document(app_path)  # if it's in a 'static' folder
    composer.append(appendix)

    # Save to buffer
    buffer = io.BytesIO()
    composer.save(buffer)

    # # Save the file to a temporary location
    # temp_file_name = "IP_Details.docx"
    # buffer.seek(0)
    # with open(temp_file_name, 'wb') as f:
    #     f.write(buffer.getbuffer())
    drive_service = create_drive_service()
    parent_folder_id = '10SOKu4DA01jm-P6U2eNwQxdTuqyk5EsS'

    # Create the media object for upload
    media = MediaIoBaseUpload(
        buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    file_metadata = {
        'name': data['naming']+'.docx',
        'parents': [parent_folder_id]
    }

    file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id'
    ).execute()

    print(f"File uploaded to Google Drive. File ID: {file.get('id')}")

    # No need to create or delete a temporary file
    return send_file(io.BytesIO(buffer.getvalue()), as_attachment=True, download_name="IP_Details.docx")

if __name__ == '__main__':
    app.run(debug=True)